#!/usr/bin/env python3
"""
generate_agenda.py - Generate IA TMC meeting agenda docx and pdf from JSON data.
Usage: python generate_agenda.py <input_json> <output_dir>
"""

import json
import sys
import os
import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _get_rpr(para):
    """Return a deepcopy of the first run's rPr element, or None."""
    for run in para.runs:
        rpr = run._r.find(qn('w:rPr'))
        if rpr is not None:
            return deepcopy(rpr)
    return None


def set_para_text(para, text):
    """Replace paragraph content with text, preserving the first run's char format."""
    if text is None:
        text = ""
    rpr = _get_rpr(para)

    keep = {qn('w:pPr'), qn('w:bookmarkStart'), qn('w:bookmarkEnd')}
    for child in list(para._p):
        if child.tag not in keep:
            para._p.remove(child)

    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    para._p.append(r)


def set_para_lines(para, lines):
    """Replace paragraph content with multiple lines joined by <w:br/>, preserving format."""
    if not lines:
        set_para_text(para, "")
        return
    rpr = _get_rpr(para)

    keep = {qn('w:pPr'), qn('w:bookmarkStart'), qn('w:bookmarkEnd')}
    for child in list(para._p):
        if child.tag not in keep:
            para._p.remove(child)

    for i, line in enumerate(lines):
        r = OxmlElement('w:r')
        if rpr is not None:
            r.append(deepcopy(rpr))
        t = OxmlElement('w:t')
        t.text = line
        if line and (line[0] == ' ' or line[-1] == ' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        para._p.append(r)
        if i < len(lines) - 1:
            br_r = OxmlElement('w:r')
            if rpr is not None:
                br_r.append(deepcopy(rpr))
            br_r.append(OxmlElement('w:br'))
            para._p.append(br_r)


def set_cell_text(cell, text, para_index=0):
    """Set text of paragraph para_index in cell."""
    paras = cell.paragraphs
    if para_index < len(paras):
        set_para_text(paras[para_index], text)


def clone_row(table, src_idx):
    """Return a deep-copied TR element of the row at src_idx."""
    return deepcopy(table.rows[src_idx]._tr)


def insert_tr_after(table, after_idx, new_tr):
    """Insert new_tr immediately after the row at after_idx."""
    table.rows[after_idx]._tr.addnext(new_tr)


def delete_row(table, idx):
    """Delete the row at idx."""
    table._tbl.remove(table.rows[idx]._tr)


def find_row(table, keyword, col=2):
    """Return index of first row where cells[col].text contains keyword, else -1."""
    for i, row in enumerate(table.rows):
        cells = row.cells
        if len(cells) > col and keyword in cells[col].text:
            return i
    return -1


# ---------------------------------------------------------------------------
# Main generation function
# ---------------------------------------------------------------------------

def generate(data: dict, template_path: str, output_docx: str, output_pdf: str = None):
    doc = Document(template_path)

    speeches      = data.get('speeches', [])
    tme           = data.get('tme', '')
    timer         = data.get('timer', '')
    ah_counter    = data.get('ah_counter', '')
    vote_counter  = data.get('vote_counter', '')
    president     = data.get('president', '')
    n_speeches    = len(speeches)

    # ---- Body paragraphs -----------------------------------------------
    for para in doc.paragraphs:
        txt = para.text
        if 'Regular meeting' in txt and 'TMC' in txt:
            new_txt = re.sub(
                r'\d+\w+\s+Regular meeting',
                f"{data.get('meeting_number', '')} Regular meeting",
                txt
            )
            set_para_text(para, new_txt)
        elif txt.startswith('Theme:'):
            set_para_text(para, f"Theme: {data.get('theme', '')}")
        elif txt.startswith('Theme question:'):
            set_para_text(para, f"Theme question: {data.get('theme_question', '')}")
        elif 'Our next meeting' in txt:
            next_meetings = data.get('next_meetings', [])
            # Fall back to legacy single-entry fields for backwards compat
            if not next_meetings:
                t = data.get('next_meeting_type', '')
                d = data.get('next_meeting_date', '')
                next_meetings = [{'type': t, 'date': d}] if t or d else []

            if next_meetings:
                ordinals = ["Our next meeting", "The following meeting", "Additionally"]
                lines = []
                for i, nm in enumerate(next_meetings):
                    prefix = ordinals[i] if i < len(ordinals) else "Also"
                    lines.append(
                        f"{prefix} is {nm.get('type','')} "
                        f"and will be held on {nm.get('date','')}."
                    )
                set_para_lines(para, lines)

    # ---- Table 0: Date / Location --------------------------------------
    tbl0 = doc.tables[0]
    date_para = tbl0.rows[0].cells[0].paragraphs[0]
    old = date_para.text
    new = re.sub(r'[A-Z][a-z]+\. \d+\w+ \([A-Za-z]+\)', data.get('date_display', ''), old)
    set_para_text(date_para, new if new != old else old.replace(
        re.search(r'[A-Z][a-z]+\. \d+\w+ \([A-Za-z]+\)', old).group()
        if re.search(r'[A-Z][a-z]+\. \d+\w+ \([A-Za-z]+\)', old) else '',
        data.get('date_display', '')
    ))

    # ---- Table 1: Main Agenda ------------------------------------------
    tbl1 = doc.tables[1]

    # Row 1 – Call to Order
    set_cell_text(tbl1.rows[1].cells[3], f"SAA:{data.get('saa', '')}", 0)
    set_cell_text(tbl1.rows[1].cells[3], f"President: {president}", 1)

    # Row 2 – TME intro
    set_cell_text(tbl1.rows[2].cells[3], f"TME: {tme}", 0)

    # Nested table: Timer / Ah Counter / Vote Counter
    nested_tbls = tbl1.rows[2].cells[2].tables
    if nested_tbls:
        nt = nested_tbls[0]
        set_cell_text(nt.rows[0].cells[1], timer)
        set_cell_text(nt.rows[1].cells[1], ah_counter)
        set_cell_text(nt.rows[2].cells[1], vote_counter)

    # Row 3 – Variety Session
    set_cell_text(tbl1.rows[3].cells[3], data.get('variety_session', ''))

    # Row 4 – Prepared Speech Session header
    set_cell_text(tbl1.rows[4].cells[3], f"TME: {tme}")

    # --- Adjust speech rows (template has 2 speeches = rows 5,6,7,8) ---
    timer1_idx = find_row(tbl1, '1st Timer')
    template_n  = (timer1_idx - 5) // 2   # should be 2

    if n_speeches < template_n:
        for i in range(template_n - 1, n_speeches - 1, -1):
            delete_row(tbl1, 5 + i * 2 + 1)   # participant sharing
            delete_row(tbl1, 5 + i * 2)         # speech row
    elif n_speeches > template_n:
        for _ in range(n_speeches - template_n):
            timer1_idx = find_row(tbl1, '1st Timer')
            insert_tr_after(tbl1, timer1_idx - 1, clone_row(tbl1, 6))  # participant
            insert_tr_after(tbl1, timer1_idx - 1, clone_row(tbl1, 5))  # speech

    # Fill speech rows
    for i, sp in enumerate(speeches):
        row = tbl1.rows[5 + i * 2]
        set_cell_text(row.cells[2],
                      f"{sp.get('pathway_abbr','')} {sp.get('level_project','')}: {sp.get('title','')}")
        set_cell_text(row.cells[3], sp.get('speaker', ''))

    # 1st Timer / Ah Counter Report
    timer1_idx = find_row(tbl1, '1st Timer')
    if timer1_idx >= 0:
        set_cell_text(tbl1.rows[timer1_idx].cells[3], f"Timer: {timer}", 0)
        set_cell_text(tbl1.rows[timer1_idx].cells[3], f"Ah Counter: {ah_counter}", 1)

    # Table Topics
    ttm_idx = find_row(tbl1, 'Table topic session')
    if ttm_idx >= 0:
        set_cell_text(tbl1.rows[ttm_idx].cells[3], data.get('table_topics_master', ''))

    # Evaluation Session header
    eval_idx = find_row(tbl1, 'Evaluation Session')
    if eval_idx >= 0:
        set_cell_text(tbl1.rows[eval_idx].cells[3], f"TME: {tme}")

    # Adjust evaluator rows
    timer2_idx = find_row(tbl1, '2nd Timer')
    cur_eval   = timer2_idx - eval_idx - 1

    if n_speeches < cur_eval:
        for i in range(cur_eval - 1, n_speeches - 1, -1):
            delete_row(tbl1, eval_idx + 1 + i)
    elif n_speeches > cur_eval:
        for _ in range(n_speeches - cur_eval):
            timer2_idx = find_row(tbl1, '2nd Timer')
            insert_tr_after(tbl1, timer2_idx - 1, clone_row(tbl1, eval_idx + 1))

    # Fill evaluator rows
    for i, sp in enumerate(speeches):
        lp        = sp.get('level_project', '')
        proj_name = sp.get('project_name', '')
        evaluator = sp.get('evaluator', '')
        suffix    = sp.get('evaluator_suffix', '')
        row = tbl1.rows[eval_idx + 1 + i]
        set_cell_text(row.cells[2], f"{lp}: {proj_name}" if proj_name else lp)
        ie_str = f"IE: {evaluator}" + (f" {suffix}" if suffix else "")
        set_cell_text(row.cells[3], ie_str)

    # 2nd Timer / Ah Counter
    timer2_idx = find_row(tbl1, '2nd Timer')
    if timer2_idx >= 0:
        set_cell_text(tbl1.rows[timer2_idx].cells[3], f"Timer: {timer}", 0)
        set_cell_text(tbl1.rows[timer2_idx].cells[3], f"Ah Counter: {ah_counter}", 1)

    # Language Evaluation / Vote Time / Small Talk
    le_idx = find_row(tbl1, 'Language Evaluation')
    if le_idx >= 0:
        set_cell_text(tbl1.rows[le_idx].cells[3], data.get('language_evaluator', ''))

    vt_idx = find_row(tbl1, 'Vote Time')
    if vt_idx >= 0:
        set_cell_text(tbl1.rows[vt_idx].cells[3], f"TME: {tme}")

    st_idx = find_row(tbl1, 'Small Talk')
    if st_idx >= 0:
        set_cell_text(tbl1.rows[st_idx].cells[3], f"TME: {tme}")

    # Feedback / Awards / Closing
    fb_idx = find_row(tbl1, 'Feedback')
    if fb_idx >= 0:
        set_cell_text(tbl1.rows[fb_idx].cells[3], f"Vote counter: {vote_counter}", 0)
        set_cell_text(tbl1.rows[fb_idx].cells[3], f"President: {president}", 1)

    # Meeting Adjournment
    adj_idx = find_row(tbl1, 'Meeting Adjournment')
    if adj_idx >= 0:
        set_cell_text(tbl1.rows[adj_idx].cells[3], f"President: {president}")

    # ---- Table 2: Speech Objectives ------------------------------------
    tbl2 = doc.tables[2]
    ROWS_PER = 3   # heading + project name + purpose
    tmpl_n2  = len(tbl2.rows) // ROWS_PER   # should be 2

    if n_speeches < tmpl_n2:
        for i in range(tmpl_n2 - 1, n_speeches - 1, -1):
            for j in range(ROWS_PER - 1, -1, -1):
                delete_row(tbl2, i * ROWS_PER + j)
    elif n_speeches > tmpl_n2:
        for _ in range(n_speeches - tmpl_n2):
            for j in range(ROWS_PER):
                tbl2._tbl.append(clone_row(tbl2, j))

    for i, sp in enumerate(speeches):
        b = i * ROWS_PER
        set_cell_text(tbl2.rows[b    ].cells[0],
                      f"{sp.get('pathway_full','')} {sp.get('level_project_full','')}:")
        set_cell_text(tbl2.rows[b + 1].cells[0], sp.get('project_name', ''))
        set_cell_text(tbl2.rows[b + 2].cells[0], sp.get('purpose', ''))

    # ---- Table 3: Officer Directory ------------------------------------
    tbl3     = doc.tables[3]
    officers = data.get('officers', [])
    for i, officer in enumerate(officers):
        row_idx = i + 1  # skip header row 0
        if row_idx >= len(tbl3.rows):
            break
        row = tbl3.rows[row_idx]
        set_cell_text(row.cells[1], officer.get('name',  ''))
        set_cell_text(row.cells[2], officer.get('email', ''))
        set_cell_text(row.cells[3], officer.get('ext',   ''))

    # ---- Save ----------------------------------------------------------
    os.makedirs(os.path.dirname(output_docx) if os.path.dirname(output_docx) else '.', exist_ok=True)
    doc.save(output_docx)

    if output_pdf:
        from docx2pdf import convert
        convert(output_docx, output_pdf)

    return output_docx, output_pdf


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: generate_agenda.py <input.json> <output_dir>", file=sys.stderr)
        sys.exit(1)

    json_path  = sys.argv[1]
    output_dir = sys.argv[2]

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    script_dir    = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.normpath(os.path.join(script_dir, '..', 'templates', 'agenda_template.docx'))

    meeting_num = data.get('meeting_number', 'meeting')
    base_name   = f"IA_TMC_{meeting_num}_Regular_Meeting"
    out_docx    = os.path.join(output_dir, base_name + '.docx')
    out_pdf     = os.path.join(output_dir, base_name + '.pdf')

    try:
        generate(data, template_path, out_docx, out_pdf)
        # Qt reads this line to get output paths
        print(f"SUCCESS|{out_docx}|{out_pdf}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"ERROR|{e}", file=sys.stderr)
        sys.exit(1)
